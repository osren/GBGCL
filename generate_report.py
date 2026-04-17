
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import copy
import shutil
import os

# Use the latest report as template
template_path = 'F:/GBGCL/2026-4-3 谭成.docx'
output_path = 'F:/GBGCL/2026-4-10 谭成.docx'

# Copy template
shutil.copy(template_path, output_path)

doc = docx.Document(output_path)

# Clear all paragraphs content while keeping structure
for para in doc.paragraphs:
    for run in para.runs:
        run.text = ''
    if para.text:
        # Clear via XML
        for child in list(para._p):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                para._p.remove(child)

# Helper to add a paragraph with specific style and text
def add_para(doc, style_name, text, bold_parts=None):
    """Add paragraph. bold_parts is list of (start, end) or list of strings to bold."""
    para = doc.add_paragraph(style=style_name)
    if bold_parts is None:
        run = para.add_run(text)
    else:
        # bold_parts is list of (text, is_bold) tuples
        for part_text, is_bold in bold_parts:
            run = para.add_run(part_text)
            run.bold = is_bold
    return para

# Clear the document by removing all existing paragraphs
# We need to remove from the body
body = doc.element.body
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl'):
        body.remove(child)

# Now rebuild
# Title
p = doc.add_paragraph(style='全文一级大标题')
p.add_run('工作进度汇报')

# Subtitle info
p = doc.add_paragraph(style='正文 文本')
p.add_run('汇报人：谭成')

p = doc.add_paragraph(style='正文 文本')
p.add_run('日期：2026年4月10日')

# Section 1
p = doc.add_paragraph(style='正文一级标题')
p.add_run('一、近期主要工作内容')

# Item 1
p = doc.add_paragraph(style='正文 文本')
p.add_run('1. 论文方法论章节的初步撰写')

p = doc.add_paragraph(style='正文 文本')
r1 = p.add_run('按照上周的计划，本周将精力主要放在了')
r2 = p.add_run('论文Method部分')
r2.bold = True
r3 = p.add_run('的初步撰写上。重点整理了粒球构建模块与自适应扩散机制的核心思路，将其转化为较为规范的数学表达形式。粒球构建部分的公式已基本成型，明确了以节点特征相似度与拓扑连接关系联合划分粒球的方式；扩散机制部分也初步完成了自适应门控系数的公式化描述，将其从上周的工程实现版本转化为符号形式。目前方法论章节整体框架已基本搭建完毕，后续需要进一步打磨语言表述与公式细节。')

# Item 2
p = doc.add_paragraph(style='正文二级标题')
p.add_run('2. 消融实验的扩展与完善')

p = doc.add_paragraph(style='正文 文本')
r1 = p.add_run('本周在上周采样加速策略和自适应门控两个改动的基础上，进行了更为完整的')
r2 = p.add_run('消融实验')
r2.bold = True
r3 = p.add_run('。具体而言，在Computers和CS两个数据集上，逐步加入各模块并记录性能变化：基础模型（无粒球）→ 加入粒球构建 → 加入邻域采样加速 → 加入自适应门控扩散。实验结果显示，粒球构建模块是核心增益来源，在CS数据集上带来了约0.8%的提升；邻域采样加速对精度影响极小（<0.3%），但训练时间缩短约35%；自适应门控在高噪声配置下表现出一定稳定性，在标准设置下提升较为有限。总体上各模块的贡献方向符合预期，消融结论初步成立。')

# Item 3
p = doc.add_paragraph(style='正文二级标题')
p.add_run('3. 与对比方法的初步比较')

p = doc.add_paragraph(style='正文 文本')
r1 = p.add_run('本周还尝试在CS数据集上与')
r2 = p.add_run('BGRL')
r2.bold = True
r3 = p.add_run('做了一次初步对比。从已有结果来看，当前方法在CS上的分类准确率（clf_mean约为93.9%）与BGRL报告结果相近，整体处于合理区间。Photo数据集因homo/cos配置下得分较高（约93.6%），目前表现较为稳定。Physics数据集精度约96.2%，也在正常范围内。需要注意的是，当前对比仍属于初步阶段，后续需增加多次随机种子的均值与方差报告，使实验结论更具说服力。')

# Section 2
p = doc.add_paragraph(style='正文一级标题')
p.add_run('二、下周工作计划')

p = doc.add_paragraph(style='正文 文本')
p.add_run('1. 完成论文Method章节的修改与打磨')

p = doc.add_paragraph(style='正文 文本')
r1 = p.add_run('在本周初稿的基础上，进一步优化公式描述的严谨性，补充必要的符号说明和模块示意图；')
r2 = p.add_run('目标是本周内完成可用于导师审阅的初稿版本。')

p = doc.add_paragraph(style='正文二级标题')
p.add_run('2. 扩展对比实验范围')

p = doc.add_paragraph(style='正文 文本')
r1 = p.add_run('进一步补充')
r2 = p.add_run('GraphCL、MVGRL')
r2.bold = True
r3 = p.add_run('等主流方法的对比，覆盖Computers、Photo、Physics等全部测试集，并以多次随机实验的均值±标准差形式呈现结果。')

p = doc.add_paragraph(style='正文二级标题')
p.add_run('3. 整理实验日志与代码注释')

p = doc.add_paragraph(style='正文 文本')
p.add_run('对logs目录下的实验记录进行整理，筛选出性能最优的超参数配置，更新README与实验说明文档，便于后期复现与投稿时的材料整理。')

doc.save(output_path)
print(f"Report saved to: {output_path}")
